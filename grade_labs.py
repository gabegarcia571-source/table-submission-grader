import os
import re
import pdfplumber
import pandas as pd
from docx import Document

# ── Configuration ─────────────────────────────────────────────────────────────

SUBMISSIONS_FOLDER = "submissions"
TOTAL_QUESTIONS    = 10
MIN_ANSWER_LENGTH  = 3

# File types that are graded (docx, pdf, plain-text variants)
GRADED_EXTENSIONS = {".docx", ".pdf", ".txt", ".pages", ".rtf", ".odt"}

# File types that are silently skipped (student uploaded extra files)
SKIPPED_EXTENSIONS = {".xlsx", ".xls", ".xlsm", ".csv", ".pptx", ".ppt", ".zip"}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _count_from_table_rows(all_rows):
    """
    Given a list of rows (each a list of cell strings), look for a header row
    that contains 'question' and 'answer', then count non-empty answer cells
    in subsequent rows. Returns count or None if no matching table found.
    """
    for i, row in enumerate(all_rows):
        cells = [c.strip().lower() for c in row]
        if any("question" in c for c in cells) and any("answer" in c for c in cells):
            answer_col = next((j for j, c in enumerate(cells) if "answer" in c), 1)
            answered = 0
            for data_row in all_rows[i + 1:]:
                if len(data_row) > answer_col:
                    question_text = data_row[0].strip()
                    answer_text   = data_row[answer_col].strip()
                    if (
                        len(question_text) > MIN_ANSWER_LENGTH
                        and "optional" not in question_text.lower()
                        and len(answer_text) > MIN_ANSWER_LENGTH
                    ):
                        answered += 1
            return answered
    return None


def _count_from_numbered_text(text):
    """
    Fallback: scan plain text for numbered answers like '1.', '2.', etc.
    """
    answered = 0
    for q in range(1, TOTAL_QUESTIONS + 1):
        pattern = rf"{q}\.?(.*?)(?=\n\d+\.|\Z)"
        match = re.search(pattern, text, re.DOTALL)
        if match and len(match.group(1).strip()) > MIN_ANSWER_LENGTH:
            answered += 1
    return answered


# ── Per-format extractors ─────────────────────────────────────────────────────

def count_answered_from_docx(path):
    """
    Handles:
      1. Table-based: two-column table with Question | Answer header
      2. Numbered text paragraphs: '1. answer text'
    """
    doc = Document(path)

    # Format 1 — table
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue
        all_rows = [[c.text for c in row.cells] for row in rows]
        result = _count_from_table_rows(all_rows)
        if result is not None:
            return result

    # Format 2 — numbered paragraphs
    text_parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    return _count_from_numbered_text("\n".join(text_parts))


def count_answered_from_pdf(path):
    """
    Handles:
      1. Table-based PDFs (most common for this assignment)
      2. Plain-text PDFs with numbered answers
    """
    all_rows  = []
    full_text = ""

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for row in table:
                    all_rows.append([cell or "" for cell in row])
            t = page.extract_text()
            if t:
                full_text += t

    result = _count_from_table_rows(all_rows)
    if result is not None:
        return result

    return _count_from_numbered_text(full_text)


def count_answered_from_text(path):
    """
    Handles plain-text formats: .txt, .rtf, .pages, .odt
    Reads as UTF-8 with fallback encoding, then applies numbered-text logic.
    """
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    except Exception:
        return 0

    # Strip RTF control words so the regex can find numbered answers
    if path.lower().endswith(".rtf"):
        text = re.sub(r"\\[a-z]+\d*\s?|\{|\}", " ", text)

    return _count_from_numbered_text(text)


# ── Dispatcher ────────────────────────────────────────────────────────────────

def count_answered_questions(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            return count_answered_from_docx(path)
        if ext == ".pdf":
            return count_answered_from_pdf(path)
        if ext in {".txt", ".rtf", ".pages", ".odt"}:
            return count_answered_from_text(path)
    except Exception as e:
        print(f"  [WARN] Could not parse {os.path.basename(path)}: {e}")
    return 0


# ── Grading logic ─────────────────────────────────────────────────────────────

def grade(answered):
    """
    Rubric:
      1 point  — answers provided for first 10 questions (wrong answers OK)
      0 points — no submission OR obviously incomplete
    """
    if answered >= 8:
        return 1, "Complete"
    elif answered > 0:
        return 0, "Incomplete"
    else:
        return 0, "No Submission / Empty"


# ── Main ──────────────────────────────────────────────────────────────────────

results = []
skipped = []

for file in sorted(os.listdir(SUBMISSIONS_FOLDER)):
    path = os.path.join(SUBMISSIONS_FOLDER, file)
    if not os.path.isfile(path):
        continue

    ext = os.path.splitext(file)[1].lower()

    # Silently skip extra file types (Excel, PowerPoint, etc.)
    if ext in SKIPPED_EXTENSIONS:
        skipped.append(file)
        continue

    # Unsupported but not explicitly skipped — record as ungradeable
    if ext not in GRADED_EXTENSIONS:
        results.append({
            "Student File":       file,
            "File Type":          ext or "(no extension)",
            "Questions Answered": 0,
            "Points (out of 1)":  0,
            "Status":             "Unsupported File Type",
        })
        continue

    answered       = count_answered_questions(path)
    points, status = grade(answered)

    results.append({
        "Student File":       file,
        "File Type":          ext,
        "Questions Answered": answered,
        "Points (out of 1)":  points,
        "Status":             status,
    })

# ── Output ────────────────────────────────────────────────────────────────────

df = pd.DataFrame(results)
df.to_csv("grading_results.csv", index=False)

total      = len(results)
complete   = (df["Points (out of 1)"] == 1).sum()
incomplete = (df["Status"] == "Incomplete").sum()
empty      = (df["Status"].isin(["No Submission / Empty", "Unsupported File Type"])).sum()

print(f"\nGrading complete — {total} submission(s) processed.")
print(f"  ✓ Complete (1 pt):   {complete}")
print(f"  ✗ Incomplete (0 pt): {incomplete}")
print(f"  ✗ Empty/No Sub (0):  {empty}")
if skipped:
    print(f"\n  Skipped {len(skipped)} non-submission file(s): {', '.join(skipped)}")
print(f"\nResults saved to grading_results.csv")